from __future__ import annotations

import hashlib
import io
from pathlib import Path
from typing import Iterable, List, Tuple

import pdfplumber
import pytesseract
from docx import Document as DocxDocument
from PIL import Image, UnidentifiedImageError

from models.schemas import IngestError, IngestedDocument


def _sha256(raw_bytes: bytes) -> str:
    return f"sha256:{hashlib.sha256(raw_bytes).hexdigest()}"


def _detect_format(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    return {
        ".pdf": "PDF",
        ".docx": "DOCX",
        ".jpg": "IMAGE",
        ".jpeg": "IMAGE",
        ".png": "IMAGE",
        ".txt": "TXT",
    }.get(extension, "UNSUPPORTED")


def _read_pdf(raw_bytes: bytes) -> Tuple[str, int, bool]:
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        pages = pdf.pages
        text_parts = [(page.extract_text() or "").strip() for page in pages]
        full_text = "\n\n".join(part for part in text_parts if part)
        if full_text:
            return full_text, max(1, len(pages)), False
    try:
        image = Image.open(io.BytesIO(raw_bytes))
        return pytesseract.image_to_string(image), 1, True
    except pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError("OCR engine not installed. Install tesseract to process scanned PDFs and images.") from exc
    except Exception:
        return "", 1, True


def _read_docx(raw_bytes: bytes) -> Tuple[str, int, bool]:
    doc = DocxDocument(io.BytesIO(raw_bytes))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
    return text, 1, False


def _read_image(raw_bytes: bytes) -> Tuple[str, int, bool]:
    try:
        image = Image.open(io.BytesIO(raw_bytes))
    except UnidentifiedImageError as exc:
        raise ValueError("INGEST_FAILED") from exc
    try:
        return pytesseract.image_to_string(image), 1, True
    except pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError("OCR engine not installed. Install tesseract to process scanned PDFs and images.") from exc


def _check_pdf_protection(raw_bytes: bytes) -> None:
    try:
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            if getattr(pdf.doc, "is_encrypted", False):
                raise ValueError("PASSWORD_PROTECTED")
    except Exception as exc:
        if isinstance(exc, ValueError) and str(exc) == "PASSWORD_PROTECTED":
            raise


def ingest_document(filename: str, raw_bytes: bytes, doc_id: str) -> IngestedDocument:
    doc_format = _detect_format(filename)
    if doc_format == "UNSUPPORTED":
        raise ValueError("UNSUPPORTED_FORMAT")
    if doc_format == "PDF":
        _check_pdf_protection(raw_bytes)
        raw_text, page_count, ocr_used = _read_pdf(raw_bytes)
        normalized_format = "PDF_OCR" if ocr_used else "PDF_TEXT"
    elif doc_format == "DOCX":
        raw_text, page_count, ocr_used = _read_docx(raw_bytes)
        normalized_format = "DOCX"
    elif doc_format == "IMAGE":
        raw_text, page_count, ocr_used = _read_image(raw_bytes)
        normalized_format = "OCR_IMAGE"
    else:
        raw_text = raw_bytes.decode("utf-8", errors="ignore")
        page_count = 1
        ocr_used = False
        normalized_format = "EMAIL_TEXT"

    return IngestedDocument(
        doc_id=doc_id,
        filename=filename,
        file_hash=_sha256(raw_bytes),
        raw_text=raw_text,
        page_count=page_count,
        ocr_used=ocr_used,
        format=normalized_format,
    )


def ingest_files(files: Iterable[object]) -> Tuple[List[IngestedDocument], List[IngestError]]:
    documents: List[IngestedDocument] = []
    errors: List[IngestError] = []
    for index, upload in enumerate(files, start=1):
        filename = getattr(upload, "filename", f"file-{index}")
        raw_bytes = upload.file.read() if hasattr(upload, "file") else bytes()
        try:
            documents.append(ingest_document(filename, raw_bytes, f"DOC-{index:03d}"))
        except ValueError as exc:
            errors.append(IngestError(filename=filename, error=str(exc), detail=None))
        except RuntimeError as exc:
            errors.append(IngestError(filename=filename, error="INGEST_FAILED", detail=str(exc)))
        except Exception as exc:
            errors.append(IngestError(filename=filename, error="INGEST_FAILED", detail=str(exc)))
    return documents, errors
